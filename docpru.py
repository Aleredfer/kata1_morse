from docx import Document
from docx.shared import Pt
from time import strftime, gmtime

document = Document()

document.add_heading('Telegrama', 0)

hoy = strftime("%d - %b - %Y", gmtime())   # te da el día, fecha y año.

fecha = document.add_paragraph(hoy)
fecha.alignment = 2

de = document.add_paragraph("")
de.add_run('from: ').bold = True
de.add_run('María Domingo')

para = document.add_paragraph("")
para.add_run('To: ').bolt = True
para.add_run('Ramón Maldonado')

document.add_heading('Mensaje', level=1)

table = document.add_table(rows=2, cols=1)
table.style = 'LightShading'

table.style.font.name ='Copurier'
table.style.font.size = Pt(12)
# es lo mismo:
# style = table.sytle
# font = style.font
# font.name ='Courier'
# font.size = Pt(12)


celda_morse = table.rows[0].cells[0]
celda_morse.text = '.-.-.'
celda_plano = table.rows[1].cells[0]
celda_plano.text = 'Hola'


document.save('demo2.docx')